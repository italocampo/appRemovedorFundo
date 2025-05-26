# Como ler e adicionar um parágrafo em um arquivo Word

import docx

# Caminho do arquivo
caminho_arquivo = r"C:\Users\Windows OS\OneDrive\Área de Trabalho\Faculdade\3° sem ADS\Phyton3\Phyton\conteudo\Word.docx"
#Esse 'r' é só para garantir que as barras invertidas sejam interpretadas corretamente.

# Abrir o documento apenas uma vez
doc = docx.Document(caminho_arquivo)

novo_doc = docx.Document()

#Excluir palavra

# remover_palavra = "!"
# for paragrafo in doc.paragraphs:
#     if remover_palavra in paragrafo.text:
#         paragrafo.text = paragrafo.text.replace(remover_palavra, "")

# #Filtrar uma palavra

# for paragrafo in doc.paragraphs:
#if "Phyton" in paragrafo.text or "python" in paragrafo.text:
#         novo_doc.add_paragraph(paragrafo.text)

# caminho_novo_arquivo = r"C:\Users\Windows OS\OneDrive\Área de Trabalho\Faculdade\3° sem ADS\Phyton3\Phyton\conteudo\Word.docx"
# novo_doc.save(caminho_novo_arquivo)
# print("\nFiltro realizado com sucesso! ")

# Adicionar um novo parágrafo
# doc.add_paragraph("Fim!")

# Salvar no mesmo caminho
doc.save(caminho_arquivo)

# Exibir todos os parágrafos do documento
for paragrafo in doc.paragraphs:
    print(paragrafo.text)
